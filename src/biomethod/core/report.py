"""Report generation for BioMethod."""

from pathlib import Path
from typing import Any
import io

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches

from biomethod.core.models import (
    AnalysisResult,
    ReproducibilityReport,
    ReproducibilityIssue,
)
from biomethod.generators.prose import ProseGenerator
from biomethod.generators.citations import CitationFormatter


def generate_methods(
    analysis: AnalysisResult,
    style: str = "generic",
    output_format: str = "markdown",
    include_citations: bool = True,
    include_supplementary: bool = True,
) -> "MethodsReport":
    """Generate a methods section from analysis results.

    Args:
        analysis: The analysis result
        style: Journal style (generic, nature, bioinformatics, plos)
        output_format: Output format (markdown, latex, docx, html)
        include_citations: Whether to include citations
        include_supplementary: Whether to generate supplementary tables

    Returns:
        MethodsReport object
    """
    return MethodsReport(
        analysis=analysis,
        style=style,
        output_format=output_format,
        include_citations=include_citations,
        include_supplementary=include_supplementary,
    )


def reproducibility_check(analysis: AnalysisResult) -> ReproducibilityReport:
    """Check reproducibility of an analysis.

    Args:
        analysis: The analysis result to check

    Returns:
        ReproducibilityReport with issues and score
    """
    report = ReproducibilityReport()

    # Sandve's 10 rules checklist
    checklist = {
        "All software versions recorded": True,
        "Random seeds specified": True,
        "No hardcoded paths": True,
        "Environment specification present": True,
        "Input data sources documented": True,
        "All parameters recorded": True,
        "Workflow documented": True,
        "Containerization used": False,
        "Version control used": False,
        "Public data repositories used": False,
    }

    # Check for version information
    tools_without_version = [t for t in analysis.tools if not t.version]
    if tools_without_version:
        checklist["All software versions recorded"] = False
        for tool in tools_without_version:
            report.issues.append(
                ReproducibilityIssue(
                    severity="warning",
                    category="version",
                    message=f"Tool '{tool.name}' has no version specified",
                    source_file=tool.source_file,
                    line_number=tool.line_number,
                    suggestion="Specify the exact version used in your analysis",
                )
            )

    # Check for random seeds
    stochastic_tools = [
        t for t in analysis.tools
        if t.category in ["single-cell", "differential-expression", "alignment"]
    ]
    for tool in stochastic_tools:
        seed_params = ["seed", "random_seed", "random_state", "-s", "--seed"]
        has_seed = any(
            param.lower() in [s.lower() for s in seed_params]
            for param in tool.parameters.keys()
        )
        if not has_seed and tool.category != "alignment":
            checklist["Random seeds specified"] = False
            report.issues.append(
                ReproducibilityIssue(
                    severity="info",
                    category="seed",
                    message=f"Tool '{tool.name}' may benefit from a random seed",
                    source_file=tool.source_file,
                    suggestion="Consider setting a random seed for reproducibility",
                )
            )

    # Check for hardcoded paths
    path_patterns = ["/home/", "/Users/", "C:\\Users", "/tmp/", "/var/"]
    for tool in analysis.tools:
        for param, value in tool.parameters.items():
            if isinstance(value, str):
                for pattern in path_patterns:
                    if pattern in value:
                        checklist["No hardcoded paths"] = False
                        report.issues.append(
                            ReproducibilityIssue(
                                severity="warning",
                                category="path",
                                message=f"Hardcoded path in {tool.name}: {value}",
                                source_file=tool.source_file,
                                suggestion="Use relative paths or environment variables",
                            )
                        )
                        break

    # Check for environment files
    if not analysis.environment.requirements_files and not analysis.environment.environment_files:
        checklist["Environment specification present"] = False
        report.issues.append(
            ReproducibilityIssue(
                severity="warning",
                category="environment",
                message="No environment specification file found",
                suggestion="Add requirements.txt or environment.yml",
            )
        )

    # Check for containers
    if analysis.environment.containers:
        checklist["Containerization used"] = True

    # Check workflow type
    if analysis.workflow_type in ["nextflow", "snakemake"]:
        checklist["Workflow documented"] = True

    # Calculate score
    passed = sum(1 for v in checklist.values() if v)
    report.score = (passed / len(checklist)) * 100
    report.checklist = checklist

    return report


class MethodsReport:
    """Container for a generated methods section."""

    def __init__(
        self,
        analysis: AnalysisResult,
        style: str = "generic",
        output_format: str = "markdown",
        include_citations: bool = True,
        include_supplementary: bool = True,
    ):
        """Initialize the methods report.

        Args:
            analysis: The analysis result
            style: Journal style
            output_format: Output format
            include_citations: Whether to include citations
            include_supplementary: Whether to include supplementary tables
        """
        self.analysis = analysis
        self.style = style
        self.output_format = output_format
        self.include_citations = include_citations
        self.include_supplementary = include_supplementary

        self._prose_generator = ProseGenerator(style)
        self._citation_formatter = CitationFormatter()

        self._generated_text: str | None = None
        self._citations: str | None = None
        self._supplementary_data: list[dict[str, Any]] | None = None

    @property
    def text(self) -> str:
        """Get the generated methods text."""
        if self._generated_text is None:
            self._generate()
        return self._generated_text

    @property
    def citations(self) -> str:
        """Get the formatted citations."""
        if self._citations is None:
            self._citations = self._citation_formatter.format_bibliography(self.analysis)
        return self._citations

    @property
    def supplementary(self) -> list[dict[str, Any]]:
        """Get supplementary table data."""
        if self._supplementary_data is None:
            self._supplementary_data = self._prose_generator.generate_supplementary_table(
                self.analysis
            )
        return self._supplementary_data

    def _generate(self) -> None:
        """Generate the methods text."""
        self._generated_text = self._prose_generator.generate_from_template(
            self.analysis,
            include_versions=True,
            include_citations=self.include_citations,
        )

    def save(self, path: str | Path) -> None:
        """Save the methods section to a file.

        Args:
            path: Output file path
        """
        path = Path(path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            self._save_docx(path)
        elif suffix == ".html":
            self._save_html(path)
        elif suffix == ".tex":
            self._save_latex(path)
        else:
            # Default to markdown
            self._save_markdown(path)

    def _save_markdown(self, path: Path) -> None:
        """Save as Markdown."""
        with open(path, "w", encoding="utf-8") as f:
            f.write(self.text)

    def _save_docx(self, path: Path) -> None:
        """Save as DOCX."""
        doc = Document()

        # Add title
        doc.add_heading("Methods", level=1)

        # Add methods text
        for paragraph in self.text.split("\n\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # Handle headers
            if paragraph.startswith("##"):
                level = paragraph.count("#", 0, 4)
                text = paragraph.lstrip("#").strip()
                doc.add_heading(text, level=min(level, 9))
            elif paragraph.startswith("**") and paragraph.endswith("**"):
                # Bold paragraph (subheading)
                p = doc.add_paragraph()
                run = p.add_run(paragraph.strip("*"))
                run.bold = True
            else:
                # Regular paragraph
                doc.add_paragraph(paragraph)

        doc.save(path)

    def _save_html(self, path: Path) -> None:
        """Save as HTML."""
        import re

        html_content = self.text

        # Convert markdown to simple HTML
        # Headers
        html_content = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html_content, flags=re.MULTILINE)
        html_content = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html_content, flags=re.MULTILINE)
        html_content = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html_content, flags=re.MULTILINE)

        # Bold
        html_content = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html_content)

        # Paragraphs
        paragraphs = html_content.split("\n\n")
        html_paragraphs = []
        for p in paragraphs:
            p = p.strip()
            if p and not p.startswith("<h"):
                p = f"<p>{p}</p>"
            html_paragraphs.append(p)

        html_body = "\n".join(html_paragraphs)

        html_doc = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Methods Section</title>
    <style>
        body {{ font-family: serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; }}
        h1, h2, h3 {{ font-family: sans-serif; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

        with open(path, "w", encoding="utf-8") as f:
            f.write(html_doc)

    def _save_latex(self, path: Path) -> None:
        """Save as LaTeX."""
        import re

        latex_content = self.text

        # Escape special characters
        for char in ["&", "%", "$", "#", "_"]:
            latex_content = latex_content.replace(char, f"\\{char}")

        # Convert markdown to LaTeX
        latex_content = re.sub(r"^### (.+)$", r"\\subsubsection{\1}", latex_content, flags=re.MULTILINE)
        latex_content = re.sub(r"^## (.+)$", r"\\subsection{\1}", latex_content, flags=re.MULTILINE)
        latex_content = re.sub(r"^# (.+)$", r"\\section{\1}", latex_content, flags=re.MULTILINE)

        # Bold
        latex_content = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", latex_content)

        with open(path, "w", encoding="utf-8") as f:
            f.write(latex_content)

    def save_citations(self, path: str | Path) -> None:
        """Save citations to a BibTeX file.

        Args:
            path: Output file path
        """
        path = Path(path)
        with open(path, "w", encoding="utf-8") as f:
            f.write(self.citations)

    def save_supplementary(self, path: str | Path) -> None:
        """Save supplementary table to Excel or CSV.

        Args:
            path: Output file path
        """
        path = Path(path)
        df = pd.DataFrame(self.supplementary)

        if path.suffix.lower() == ".xlsx":
            df.to_excel(path, index=False, sheet_name="Software Versions")
        else:
            df.to_csv(path, index=False)

    def __str__(self) -> str:
        """Return the methods text."""
        return self.text
